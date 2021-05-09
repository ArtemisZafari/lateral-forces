from input import *

class Wind:
    
    def __init__(self, roof_pitch, floors, kzt, wind_speed, wind_exposure):
        self.roof_pitch = roof_pitch
        self.floors = floors
        self.kzt = kzt
        self.wind_speed = wind_speed
        self.wind_exposure = wind_exposure

    def __str__(self):
        return "This is a wind report object for a {0:0.0f} story house.".format(self.floors)

    def __len__(self):
        return self.floors
    
    def factors(self):
        # This function establishes the A, B, C, and D factors

        # Here we check our manually entered combinations of roof pitch and wind speed
        # in order to find the of A, B, C, and D values
        if 25 <= self.roof_pitch < 30 and self.wind_speed == 110:
            factors = [21.6, 14.8, 17.2, 11.8]

        elif 20 <= self.roof_pitch < 25 and self.wind_speed == 110:
            factors = [24.1, 3.9, 17.4, 4]

        elif 15 <= self.roof_pitch < 20 and self.wind_speed == 110:
            factors = [20, 4, 15, 8]

        elif 10 <= self.roof_pitch < 15 and self.wind_speed == 110:
            factors = [19.2, 0, 12.7, 0]

        else:
            print ('That combination of roof pitch and wind speed is not in the database.')
            factors = [9999, 9999, 9999, 9999]
            
        self.A = factors[0]
        self.B = factors[1]
        self.C = factors[2]
        self.D = factors[3]
        
    def calculate(self, floor_areas, A, B, C, D):
        # This function returns the wind load per floor in the units of kips in the form of a list

        # Here we go through each floor and multiply the areas by the factors
        floor_loads = []
        for floor in range(self.floors):
            current_floor_area = floor_areas[(floor*4):(floor*4)+4]
            current_floor_load = ((A*current_floor_area[0] + B*current_floor_area[1]
            + C*current_floor_area[2] + D*current_floor_area[3])/1000)*self.kzt*0.6
            floor_loads.append(current_floor_load)
        return floor_loads
    
    def control(self, loads_1, loads_2):
        # This function returns a list of strings for each floor
        # that identifies whether the non-min loading controls that
        # floor in the specified direction

        # Here we go through the floors and identify if the non-min loading controls
        control_list = []
        for floor in range(self.floors):
            if loads_1[floor] > loads_2[floor]:
                control_list.append('(controlling)')
            else:
                control_list.append('')
        return control_list
        
    def work(self, floor_areas, wind_loads, control_list, A, B, C, D):
        # This function takes an input of floors, floor_areas, and a direction indicator
        # and returns a single string that shows the calculations for each floor

        # Here we create a list of lists: the inner lists are composed of
        # the A, B, C, and D area values for each floor (we use the globalized floor_loads)
        organized_floor_areas = []
        for floor in range(self.floors):
            current_floor_area = floor_areas[(floor*4):(floor*4)+4]
            organized_floor_areas.append(current_floor_area)

        # Here we create a list of strings that show the work for each floor
        work_text_list = []
        for i,group in enumerate(organized_floor_areas):
            work_text = '\nΣ Fw Floor {0:0.0f}:\n( ( ({1:0.0f} x {2:0.1f}) + ({3:0.0f} x {4:0.1f}) + ({5:0.0f} x {6:0.1f}) + ({7:0.0f} x {8:0.1f}) ) / 1000 ) x {9:0.1f} Kzt x 0.6 =\n {10:>150.2f} kip\n{11:>152}'.format(self.floors-i, 
            group[0], A, group[1], B, group[2], C, group[3], D, self.kzt, wind_loads[i], control_list[i])
            work_text_list.append(work_text)

        # Here we combine the list of strings into a single string and then return the string
        work_string = ''
        for item in work_text_list:
            work_string = work_string + item
        return work_string
        
    def printer(self, x_work, y_work, xmin_work, ymin_work):
        # This function returns a string that serves as the calculations for the wind load page

        wind_exposure_table = {0.8:'A', 1.0:'B', 1.2:'C', 1.4:'D', 1.6:'F'}
        return """WIND LOADS STARTING INFORMATION\n                            
Roof pitch: {}°             Kzt: {}            Wind speed (mph): {}            Floors: {}
Design wind pressures (psf): A={} B={} C={} D={}
Minimum wind pressures (psf): Wall=16 Roof=8            Wind exposure: {} ({})
________________________________________________________________________________________________________
CALCULATIONS FOR X-X DIRECTION:\n{}{}
________________________________________________________________________________________________________
CALCULATIONS FOR Y-Y DIRECTION:\n{}{}""".format(self.roof_pitch, self.kzt, self.wind_speed, self.floors, self.A, self.B, self.C,
        self.D, wind_exposure_table[self.wind_exposure], self.wind_exposure, x_work, xmin_work, y_work, ymin_work)

# Here we calculate the wind loads in the x and y directions
wind_report = Wind(roof_pitch, floors, kzt, wind_speed, wind_exposure)
wind_report.factors()
x_wind = wind_report.calculate(floor_areas_x, wind_report.A, wind_report.B, wind_report.C, wind_report.D)
y_wind = wind_report.calculate(floor_areas_y, wind_report.A, wind_report.B, wind_report.C, wind_report.D)
xmin_wind = wind_report.calculate(floor_areas_x, min_wall, min_roof, min_wall, min_roof)
ymin_wind = wind_report.calculate(floor_areas_y, min_wall, min_roof, min_wall, min_roof)

# Here we create the control lists that will be used to determine which force controls each floor
x_control = wind_report.control(x_wind, xmin_wind)
y_control = wind_report.control(y_wind, ymin_wind)
xmin_control = wind_report.control(xmin_wind, x_wind)
ymin_control = wind_report.control(ymin_wind, y_wind)

# Here we create the strings showing the calculations
x_work = wind_report.work(floor_areas_x, x_wind, x_control, wind_report.A, wind_report.B, wind_report.C, wind_report.D)
y_work = wind_report.work(floor_areas_y, y_wind, y_control, wind_report.A, wind_report.B, wind_report.C, wind_report.D)
xmin_work = wind_report.work(floor_areas_x, xmin_wind, xmin_control, min_wall, min_roof, min_wall, min_roof)
ymin_work = wind_report.work(floor_areas_y, ymin_wind, ymin_control, min_wall, min_roof, min_wall, min_roof)

# Here we create the text showing the calculations and assemble the completed report
wind_complete = wind_report.printer(x_work, y_work, xmin_work, ymin_work)

# WIND CALCULATIONS END, QUAKE CALCULATIONS BEGIN







class Quake:
    
    def __init__(self, floors, floor_areas, floor_DL, roof_DL, wall_DL, rho, cs, quake_factor, floor_plates):
        self.floors = floors
        self.floor_areas = floor_areas
        self.floor_DL = floor_DL
        self.roof_DL = roof_DL
        self.wall_DL = wall_DL
        self.rho = rho
        self.cs = cs
        self.quake_factor = quake_factor
        self.floor_plates = floor_plates

    def __str__(self):
        return "This is a quake report object for a {0:0.0f} story house.".format(self.floors)

    def __len__(self):
        return self.floors
    
    def area(self):
        # This function takes in the floor areas and returns a single string
        # that lists the areas for each floor

        # Here we find the number of floors, then iterate through a list of the number of
        # floors and add the floor specific info at each iteration
        floors_text = ''
        for i in range(self.floors):
            floor_text = 'Floor {}: {} ft^2      '.format(self.floors-i, self.floor_areas[i])
            floors_text = floors_text + floor_text
        return floors_text
    
    def dead(self):
        # This function takes in the floor areas and returns a single string
        # that contains the "Dead Load of Structure" information

        # Here we find the number of floors, then iterate through a list of the number of
        # floors and add the floor specific info at each iteration
        floors_text = ''
        floor_weight_sum = 0
        for i in range(self.floors):
            if i == self.floors-1:
                floor_weight = (self.floor_areas[i] * (self.roof_DL + self.wall_DL))/1000
                floor_weight_sum = floor_weight_sum + floor_weight
                floor_text = 'Floor {0:0.0f} Weight: {1:0.0f} x ({2:0.0f} + {3:0.0f}) = {4:0.2f} kip\n'.format(self.floors-i, self.floor_areas[i], self.roof_DL, self.wall_DL, floor_weight)
                floors_text = floors_text + floor_text
            else:
                floor_weight = (self.floor_areas[i] * (self.floor_DL + self.wall_DL))/1000
                floor_weight_sum = floor_weight_sum + floor_weight
                floor_text = 'Floor {0:0.0f} Weight: {1:0.0f} x ({2:0.0f} + {3:0.0f}) = {4:0.2f} kip\n'.format(self.floors-i, self.floor_areas[i], self.floor_DL, self.wall_DL, floor_weight)
                floors_text = floors_text + floor_text

        total_q_load = self.cs*floor_weight_sum*self.rho*self.quake_factor
        self.total_q = total_q_load
        first_line = 'Floor Weight Total: {0:0.2f} kip\n'.format(floor_weight_sum)
        second_line = 'Vbase = Cs x Σ Weight = {0:0.2f} x {1:0.2f} = {2:0.2f} kip\n'.format(self.cs, floor_weight_sum, self.cs*floor_weight_sum)
        third_line = 'Vasd = Vbase x {0:0.1f} x ρ = {1:0.2f} kip'.format(self.quake_factor, total_q_load)
        floors_text = floors_text + first_line + second_line + third_line
        return floors_text
    
    def lateral(self):
        # This function takes in the floor areas and returns a single string
        # that contains the "Lateral Forces" information

        floor_weights = []
        weight_height = []
        for i in range(self.floors):
            if i == self.floors-1:
                floor_weight = (self.floor_areas[i] * (self.roof_DL + self.wall_DL))/1000
                weight_height.append(floor_weight*self.floor_plates[i])
                floor_weights.append(floor_weight)
            else:
                floor_weight = (self.floor_areas[i] * (self.floor_DL + self.wall_DL))/1000
                weight_height.append(floor_weight*self.floor_plates[i])
                floor_weights.append(floor_weight)
        self.floor_weights = floor_weights
        self.weight_height = weight_height

        weight_height_sum = sum(weight_height)
        floors_text = ''
        for i in range(self.floors):
            floor_text_1 = 'Floor {0:0.0f}: {1:0.2f} kip x {2:0.0f} ft = {3:0.2f}\n'.format(self.floors-i, floor_weights[i], self.floor_plates[i], weight_height[i])
            floors_text = floors_text + floor_text_1

        floor_q_load = []
        for i in range(self.floors):
            if i == self.floors-1:
                floor_q_load.append(self.total_q*(weight_height[i]/weight_height_sum))
                floor_text_2 = 'Fq Floor {0:0.0f}: {1:0.2f} x ({2:0.2f} / {3:0.2f}) = {4:0.2f} kip'.format(self.floors-i, self.total_q, weight_height[i], weight_height_sum, floor_q_load[i])
                floors_text = floors_text + floor_text_2
            else:
                floor_q_load.append(self.total_q*(weight_height[i]/weight_height_sum))
                floor_text_2 = 'Fq Floor {0:0.0f}: {1:0.2f} x ({2:0.2f} / {3:0.2f}) = {4:0.2f} kip\n'.format(self.floors-i, self.total_q, weight_height[i], weight_height_sum, floor_q_load[i])
                floors_text = floors_text + floor_text_2
        return floors_text
    
    def forces(self):
        # This function creates a list of the quake loads per floor

        weight_height_sum = sum(self.weight_height)
        floors_text = ''
        for i in range(self.floors):
            floor_text_1 = 'Floor {0:0.0f}: {1:0.2f} kip x {2:0.0f} ft = {3:0.2f}\n'.format(i+1, self.floor_weights[i], self.floor_plates[i], self.weight_height[i])
            floors_text = floors_text + floor_text_1

        floor_q_load = []
        for i in range(self.floors):
            if i == self.floors-1:
                floor_q_load.append(self.total_q*(self.weight_height[i]/weight_height_sum))
            else:
                floor_q_load.append(self.total_q*(self.weight_height[i]/weight_height_sum))
        return floor_q_load
    
    def wind_sum(self, wind_loads, min_wind_loads):
        # This function creates a list (can be in the x or y direction) that contains
        # the controlling wind loads (summed as they go down)

        # Here we go through the wind loads and make a list (in the x direction)
        # containing the controling load (based on minimum/calculated loads)
        control_wind = []
        for i in range(self.floors):
            if wind_loads[i] > min_wind_loads[i]:
                control_wind.append(wind_loads[i])
            else:
                control_wind.append(min_wind_loads[i])

        # Here we go through the list we just made and sum up the loads that
        # are above a floor. Essentially, we are carrying the loads down
        for i in range(self.floors):
            if i == 0:
                control_wind[i] = control_wind[i]
            else:
                control_wind[i] = control_wind[i] + control_wind[i-1]
        return control_wind
    
    def quake_sum(self, quake_loads):
        # Here we repeat the summing but for the quake loads
        
        quake_summed_loads = []
        for i in range(self.floors):
            quake_summed_loads.append(quake_loads[i])

        for i in range(self.floors):
            if i == 0:
                quake_summed_loads[i] = quake_summed_loads[i]
            else:
                quake_summed_loads[i] = quake_summed_loads[i] + quake_summed_loads[i-1]
        return quake_summed_loads
    
    def governing(self, xwind_sum, ywind_sum, quake_sum):
        # This function creates a string that states which loading (quake/wind)
        # controls in each floor/direction, while also showing the load value
        # that governs (the load value accounts for all loads above the floor)

        governing_loads_x = []
        for i in range(self.floors):
            if xwind_sum[i] > quake_sum[i]:
                governing_loads_x.append(['wind', xwind_sum[i]])
            else:
                governing_loads_x.append(['quake', quake_sum[i]])

        governing_loads_y = []
        for i in range(self.floors):
            if ywind_sum[i] > quake_sum[i]:
                governing_loads_y.append(['wind', ywind_sum[i]])
            else:
                governing_loads_y.append(['quake', quake_sum[i]])

        floors_text = ''
        for i in range(self.floors):
            floor_text_1 = 'Floor {0:0.0f} is governed by {1:s} ({2:0.2f} kips) in the X-X direction and {3:s} ({4:0.2f} kips) in the Y-Y direction.\n'.format(self.floors-i,
            governing_loads_x[i][0], governing_loads_x[i][1], governing_loads_y[i][0], governing_loads_y[i][1])
            floors_text = floors_text + floor_text_1
        return floors_text
    
    def printer(self, areas, dead, lateral, governing):
        # This function returns a string that serves as the calculations for the quake page

        return """QUAKE LOADS STARTING INFORMATION\n                            
Cs: {}          Rho: {}        Quake load factor: {}
Roof Dead Load: {} psf      Floor Dead Load: {} psf      Wall Dead Load: {} psf
{}
________________________________________________________________________________________________________
DEAD LOAD OF STRUCTURE\n
{}
________________________________________________________________________________________________________
LATERAL FORCES\n
{}
________________________________________________________________________________________________________
GOVERNING LOADS\n
{}""".format(self.cs, self.rho, self.quake_factor, self.roof_DL, self.floor_DL, self.wall_DL, areas, dead, lateral, governing)

# Here we create the text for the various components of the report
quake_report = Quake(floors, floor_areas, floor_DL, roof_DL, wall_DL, rho, cs, quake_factor, floor_plates)
areas = quake_report.area()
dead = quake_report.dead()
lateral = quake_report.lateral()
quake_loads = quake_report.forces()

# Here we create the summed lists of quake and wind loads, then determine which governs each floor/direction
xwind_sum = quake_report.wind_sum(x_wind, xmin_wind)
ywind_sum = quake_report.wind_sum(y_wind, ymin_wind)
quake_sum = quake_report.quake_sum(quake_loads)
governing = quake_report.governing(xwind_sum, ywind_sum, quake_sum)

# Here we create the text showing the calculations and assemble the completed report
quake_complete = quake_report.printer(areas, dead, lateral, governing)

# QUAKE CALCULATIONS END, SHEAR CALCULATIONS BEGIN







class Shear:
    
    def __init__(self, floors):
        self.floors = floors

    def __str__(self):
        return "This is a shear report object for a {0:0.0f} story house.".format(self.floors)

    def __len__(self):
        return self.floors
    
    def wind_control(self, wind, min_wind):
    
        new_list = []
        for i in range(self.floors):
            if wind[i] > min_wind[i]:
                new_list.append(wind[i])
            else:
                new_list.append(min_wind[i])
        return new_list
    
    def carry(self, load, length, overhang_dist, distances, shear_layer_lengths, offset_dist):
        # This function returns a list of location-based loads (for each shear layer) for a floor, based off a constant datum
        # Note that this function takes into account only the loads from one floor, not of any from above

        distr = (load)/length
        carry_list = []
        for i in range(len(shear_layer_lengths)):
            if i == 0:
                layer_load = distr * (overhang_dist[0] + distances[i]/2)
                distance = offset_dist[0]
                carry_list.append([layer_load, distance])
            elif i == len(shear_layer_lengths)-1:
                layer_load = (distr) * (overhang_dist[1] + distances[i-1]/2)
                distance = sum(distances) + offset_dist[0]
                carry_list.append([layer_load, distance])
            else:
                layer_load = (distr) * (distances[i-1]/2 + distances[i]/2)
                distance = sum(distances[0:i]) + offset_dist[0]
                carry_list.append([layer_load, distance])
        return carry_list
    
    def carry_list(self, load, length, overhang_dist, distances, shear_layer_lengths, offset_dist):
        # This function returns a list that contains all the location-based loads for a direction

        load_carry_list = []
        for i in range(self.floors):
            load_carry_list.append(self.carry(load[i], length[i], overhang_dist[i],
            distances[i], shear_layer_lengths[i], offset_dist[i]))
        return load_carry_list
    
    def text_maker(self, floor_above, current_dist, prev_dist, next_dist):
        # This function takes info relevant to a shear layer and returns a list with
        # two items: text of the loads being added from above and the total load being added

        text = ''
        new_load = 0
        for load, load_dist in floor_above:
            if load_dist == current_dist or (prev_dist == 'none' and load_dist < current_dist) or (next_dist == 'none' and load_dist > current_dist):
                new_load = new_load + load
                text = text+'+ {0:0.2f} kips '.format(load)
            elif load_dist < current_dist and load_dist > prev_dist:
                difference = current_dist - prev_dist
                new_load = new_load + (load * ((load_dist - prev_dist) / (difference)))
                text = text+'+ ({0:0.2f} kips x {1:0.2f} ft/{2:0.2f} ft) '.format(load, load_dist - prev_dist, difference)
            elif load_dist > current_dist and load_dist < next_dist:
                difference = next_dist - current_dist
                new_load = new_load + (load * ((next_dist - load_dist) / (difference)))
                text = text+'+ ({0:0.2f} kips x {1:0.2f} ft/{2:0.2f} ft) '.format(load, next_dist - load_dist, difference)
        return [text, new_load]
    
    def load_add(self, new_list, original_list):
        # This function takes in 2 identical lists (created by the load_carry_list() function) and 
        # edits the new_list so that all the loads are carried down. Then the function returns a
        # a list containing 2 items: the updated new_list and a list containing the text created by
        # the text_maker() function for each shear layer

        text_list = []
        for floor, loads in enumerate(original_list):
            if floor != 0:
                for i in range(len(loads)): # this goes through each shear layer in the current floor
                    if i == 0:
                        [new_text, new_load] = self.text_maker(new_list[floor-1], loads[i][1], 'none', loads[i+1][1])
                        new_list[floor][i][0] = new_list[floor][i][0] + new_load
                        text_list.append(new_text)
                    elif i == len(loads)-1:
                        [new_text, new_load] = self.text_maker(new_list[floor-1], loads[i][1], loads[i-1][1], 'none')
                        new_list[floor][i][0] = new_list[floor][i][0] + new_load
                        text_list.append(new_text)
                    else:
                        [new_text, new_load] = self.text_maker(new_list[floor-1], loads[i][1], loads[i-1][1], loads[i+1][1])
                        new_list[floor][i][0] = new_list[floor][i][0] + new_load
                        text_list.append(new_text)
        return [new_list, text_list]
    
    def wall_sum(self, shear_lengths):
        # This function will take an input of a list of shear wall lengths
        # and return a string with the sum of the numbers shown

        text = ''
        for i,num in enumerate(shear_lengths):
            if i == 0:
                new_text = str(num) + ' ft'
                text = text + new_text
            elif i == len(shear_lengths):
                new_text = '+ ' + str(num) + ' ft'
                text = text + new_text
            else:
                new_text = ' + ' + str(num) + ' ft'
                text = text + new_text
        return text

    def reduce_text(self, load, lengths_list, floor_num):
        # This function will take 3 inputs: a SW load (in lb/ft), a list of shear wall lengths, and the current
        # floor. It will return a string which contains any possible capacity reductions necesssary

        min_wall = min(lengths_list)
        height = floor_heights[self.floors - floor_num]
        ratio =  height / min_wall
        factor = 1.25 - (0.125 * (height / min_wall))

        if ratio > 3.5:
            text = 'LRP, STRAP TO WINDOW, OR REDESIGN\n\n'
            sw = 0

        elif 3.5 >= ratio > 2:
            if 0 <= load <= 240*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 240, 240*factor)
                sw = 1
            elif 240*factor < load <= 350*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 350, 350*factor)
                sw = 2
            elif 350*factor < load <= 450*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 450, 450*factor)
                sw = 3
            elif 450*factor < load <= 585*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 585, 585*factor)
                sw = 4
            elif 585*factor < load <= 700*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 700, 700*factor)
                sw = 5
            elif 700*factor < load <= 900*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 900, 900*factor)
                sw = 6
            elif 900*factor < load <= 1170*factor:
                text = '(1.25 - 0.125x({0:0.2f} ft/{1:0.2f} ft)) x {2:0.2f} lb/ft = {3:0.2f} lb/ft capacity\nper table 4.3.4 ANSI/AF&PA SDPWS\n\n'.format(height, min_wall, 1170, 1170*factor)
                sw = 7
            elif load > 1170*factor:
                text = 'LOAD IS TOO LARGE\n'
                sw = 8

        elif ratio <= 2:
            text = ''
            if 0 <= load <= 240:
                sw = 1
            elif 240 < load <= 350:
                sw = 2
            elif 350 < load <= 450:
                sw = 3
            elif 450 < load <= 585:
                sw = 4
            elif 585 < load <= 700:
                sw = 5
            elif 700 < load <= 900:
                sw = 6
            elif 900 < load <= 1170:
                sw = 7
            elif load > 1170:
                text = 'LOAD IS TOO LARGE\n'
                sw = 8
        return [text, sw]

    def SW_num(self, sw):
        # This function will take an input from reduce_text() and return
        # the SW strength required

        if sw == 1:
            text = '[SW-6'
        elif sw == 2:
            text = '[SW-4'
        elif sw == 3:
            text = '[SW-3'
        elif sw == 4:
            text = '[SW-2'
        elif sw == 5:
            text = '[(2) SW-4'
        elif sw == 6:
            text = '[(2) SW-3'
        elif sw == 7:
            text = '[(2) SW-2'
        elif sw == 0 or sw == 8:
            text = '[???'
        return text
    
    def top_floor_text(self, load, length, distances, direction_indicator, overhang_dist, shear_layer_lengths):
        # This function returns a string that displays all the shear wall calculations for the top floor

        if direction_indicator == 'x':
            direction = 'X-X'
            direction2 = 'X'
        elif direction_indicator == 'y':
            direction = 'Y-Y'
            direction2 = 'Y'

        distr = (1000*load)/length
        v_num = 1
        shear_text = ''
        for i in range(len(shear_layer_lengths)): # this goes through each shear layer in the current floor
            if i == 0:
                layer_load = (distr/1000) * (overhang_dist[0] + distances[i]/2)
                shear_load = (layer_load * 1000) / sum(shear_layer_lengths[i])
                shear_text_1 = 'V @ {0:s} {1:0.0f} = ({2:0.2f} kips / {3:0.2f} ft) x ({4:0.2f} ft + {5:0.2f} ft/2) = {6:0.2f} kips\n'.format(direction2, v_num, load, length,
                overhang_dist[0], distances[i], layer_load)
                [shear_text_3, sw]  = self.reduce_text(shear_load, shear_layer_lengths[i], self.floors)
                shear_text_2 = '{0:0.2f} kips / ({1:s}) = {2:>20.2f} lb/ft {3:>8s}]\n\n'.format(layer_load, self.wall_sum(shear_layer_lengths[i]), shear_load, self.SW_num(sw))
                shear_text = shear_text + shear_text_1 + shear_text_2 + shear_text_3
                v_num = v_num + 1
            elif i == len(shear_layer_lengths)-1:
                layer_load = (distr/1000) * (overhang_dist[1] + distances[i-1]/2)
                shear_load = (layer_load * 1000) / sum(shear_layer_lengths[i])
                shear_text_1 = 'V @ {0:s} {1:0.0f} = ({2:0.2f} kips / {3:0.2f} ft) x ({4:0.2f} ft/2 + {5:0.2f} ft) = {6:0.2f} kips\n'.format(direction2, v_num, load, length,
                distances[i-1], overhang_dist[1], layer_load)
                [shear_text_3, sw]  = self.reduce_text(shear_load, shear_layer_lengths[i], self.floors)
                shear_text_2 = '{0:0.2f} kips / ({1:s}) = {2:>20.2f} lb/ft {3:>8s}]\n\n'.format(layer_load, self.wall_sum(shear_layer_lengths[i]), shear_load, self.SW_num(sw))
                shear_text = shear_text + shear_text_1 + shear_text_2 + shear_text_3
                v_num = v_num + 1
            else:
                layer_load = (distr/1000) * (distances[i-1]/2 + distances[i]/2)
                shear_load = (layer_load * 1000) / sum(shear_layer_lengths[i])
                shear_text_1 = 'V @ {0:s} {1:0.0f} = ({2:0.2f} kips / {3:0.2f} ft) x ({4:0.2f} ft/2 + {5:0.2f} ft/2) = {6:0.2f} kips\n'.format(direction2, v_num, load, length,
                distances[i-1], distances[i], layer_load)
                [shear_text_3, sw]  = self.reduce_text(shear_load, shear_layer_lengths[i], self.floors)
                shear_text_2 = '{0:0.2f} kips / ({1:s}) = {2:>20.2f} lb/ft {3:>8s}]\n\n'.format(layer_load, self.wall_sum(shear_layer_lengths[i]), shear_load, self.SW_num(sw))
                shear_text = shear_text + shear_text_1 + shear_text_2 + shear_text_3
                v_num = v_num + 1
        return (shear_text)
    
    def other_floor_text(self, carried_loads, carried_texts, shear_layers, floor_num, load, length, distances, direction_indicator, overhang_dist, shear_layer_lengths):
        # This function returns a string that displays all the shear wall calculations for a non-top floor

        if direction_indicator == 'x':
            direction = 'X-X'
            direction2 = 'X'
        elif direction_indicator == 'y':
            direction = 'Y-Y'
            direction2 = 'Y'

        distr = (1000*load)/length
        v_num = sum(shear_layers[0:self.floors-floor_num]) + 1
        shear_text = ''
        for i in range(len(shear_layer_lengths)): # this goes through each shear layer in the current floor
            if i == 0:
                layer_load = (distr/1000) * (overhang_dist[0] + distances[i]/2)
                shear_text_1 = 'V @ {0:s} {1:0.0f} = ({2:0.2f} kips / {3:0.2f} ft) x ({4:0.2f} ft + {5:0.2f} ft/2) = {6:0.2f} kips\n'.format(direction2, v_num, load, length,
                overhang_dist[0], distances[i], layer_load)
                carry_load = carried_loads[self.floors-floor_num][i][0]
                shear_load = (carry_load * 1000) / sum(shear_layer_lengths[i])
                [shear_text_4, sw]  = self.reduce_text(shear_load, shear_layer_lengths[i], floor_num)
                carry_text = carried_texts[i+sum(shear_layers[1:self.floors-floor_num])]
                shear_text_2 = '{0:0.2f} kips {1:s} = {2:0.2f} kips\n'.format(layer_load, carry_text, carry_load)
                shear_text_3 = '{0:0.2f} kips / ({1:s}) = {2:>20.2f} lb/ft {3:>8s}]\n\n'.format(carry_load, self.wall_sum(shear_layer_lengths[i]), shear_load, self.SW_num(sw))
                shear_text = shear_text + shear_text_1 + shear_text_2 + shear_text_3 + shear_text_4
                v_num = v_num + 1
            elif i == len(shear_layer_lengths)-1:
                layer_load = (distr/1000) * (overhang_dist[1] + distances[i-1]/2)
                shear_text_1 = 'V @ {0:s} {1:0.0f} = ({2:0.2f} kips / {3:0.2f} ft) x ({4:0.2f} ft/2 + {5:0.2f} ft) = {6:0.2f} kips\n'.format(direction2, v_num, load, length,
                distances[i-1], overhang_dist[1], layer_load)
                carry_load = carried_loads[self.floors-floor_num][i][0]
                shear_load = (carry_load * 1000) / sum(shear_layer_lengths[i])
                [shear_text_4, sw]  = self.reduce_text(shear_load, shear_layer_lengths[i], floor_num)
                carry_text = carried_texts[i+sum(shear_layers[1:self.floors-floor_num])]
                shear_text_2 = '{0:0.2f} kips {1:s} = {2:0.2f} kips\n'.format(layer_load, carry_text, carry_load)
                shear_text_3 = '{0:0.2f} kips / ({1:s}) = {2:>20.2f} lb/ft {3:>8s}]\n\n'.format(carry_load, self.wall_sum(shear_layer_lengths[i]), shear_load, self.SW_num(sw))
                shear_text = shear_text + shear_text_1 + shear_text_2 + shear_text_3 + shear_text_4
                v_num = v_num + 1
            else:
                layer_load = (distr/1000) * (distances[i-1]/2 + distances[i]/2)
                shear_text_1 = 'V @ {0:s} {1:0.0f} = ({2:0.2f} kips / {3:0.2f} ft) x ({4:0.2f} ft/2 + {5:0.2f} ft/2) = {6:0.2f} kips\n'.format(direction2, v_num, load, length,
                distances[i-1], distances[i], layer_load)
                carry_load = carried_loads[self.floors-floor_num][i][0]
                shear_load = (carry_load * 1000) / sum(shear_layer_lengths[i])
                [shear_text_4, sw]  = self.reduce_text(shear_load, shear_layer_lengths[i], floor_num)
                carry_text = carried_texts[i+sum(shear_layers[1:self.floors-floor_num])]
                shear_text_2 = '{0:0.2f} kips {1:s} = {2:0.2f} kips\n'.format(layer_load, carry_text, carry_load)
                shear_text_3 = '{0:0.2f} kips / ({1:s}) = {2:>20.2f} lb/ft {3:>8s}]\n\n'.format(carry_load, self.wall_sum(shear_layer_lengths[i]), shear_load, self.SW_num(sw))
                shear_text = shear_text + shear_text_1 + shear_text_2 + shear_text_3 + shear_text_4
                v_num = v_num + 1
        return shear_text
    
    def shear_text_floor(self, carried_loads, carried_texts, shear_layers, floor_num, load, length, distances, direction_indicator, overhang_dist, shear_layer_lengths):
        # This function determines and returns either a top_floor_text or other_floor_text output

        if direction_indicator == 'x':
            direction = 'X-X'
            direction2 = 'X'
        elif direction_indicator == 'y':
            direction = 'Y-Y'
            direction2 = 'Y'
        distr = (1000*load)/length

        first_line = '________________________________________________________________________________________________________\nFw {0:s} @ Floor {1:0.0f}: {2:0.2f} kips / {3:0.2f} ft = {4:0.2f} lb/ft\n\n'.format(direction, floor_num, load,
        length, distr)
        if floor_num == self.floors:
            shear_text = self.top_floor_text(load, length, distances, direction_indicator, overhang_dist, shear_layer_lengths)
        else:
            shear_text = self.other_floor_text(carried_loads, carried_texts, shear_layers, floor_num, load, length, distances, direction_indicator, overhang_dist, shear_layer_lengths)
        return (first_line + shear_text)
    
    def shear_text_direction(self, carried_loads, carried_texts, shear_layers, load, length, distances, direction_indicator, overhang_dist, shear_layer_lengths):
        # This function returns a string that displays all the shear wall calculations for a direction

        shear_text = ''
        for i in range(self.floors):
            text = self.shear_text_floor(carried_loads, carried_texts, shear_layers, self.floors-i, load[i], length[i], distances[i], direction_indicator, overhang_dist[i], shear_layer_lengths[i])
            shear_text = shear_text + text
        return shear_text
    
    def printer(self, wind_x, wind_y, quake_x, quake_y):
        # This function returns a string that serves as the calculations for the shear page

        return """Shearwalls in X-X Direction (WIND)
{}
________________________________________________________________________________________________________
Shearwalls in Y-Y Direction (WIND)
{}
________________________________________________________________________________________________________
Shearwalls in X-X Direction (QUAKE)
{}
________________________________________________________________________________________________________
Shearwalls in Y-Y Direction (QUAKE)
{}""".format(wind_x, wind_y, quake_x, quake_y)

# Here we create lists of the largest wind forces for each floor
shear_report = Shear(floors)
x_wind = shear_report.wind_control(x_wind, xmin_wind)
y_wind = shear_report.wind_control(y_wind, ymin_wind)

# Here we create lists containing the location-based loads of each force and direction
x1_wind = shear_report.carry_list(x_wind, floor_lengths_x, overhang_dist_x, distances_x, shear_layer_lengths_x, offset_dist_x)
y1_wind = shear_report.carry_list(y_wind, floor_lengths_y, overhang_dist_y, distances_y, shear_layer_lengths_y, offset_dist_y)
x1_quake = shear_report.carry_list(quake_loads, floor_lengths_x, overhang_dist_x, distances_x, shear_layer_lengths_x, offset_dist_x)
y1_quake = shear_report.carry_list(quake_loads, floor_lengths_y, overhang_dist_y, distances_y, shear_layer_lengths_y, offset_dist_y)
x2_wind = shear_report.carry_list(x_wind, floor_lengths_x, overhang_dist_x, distances_x, shear_layer_lengths_x, offset_dist_x)
y2_wind = shear_report.carry_list(y_wind, floor_lengths_y, overhang_dist_y, distances_y, shear_layer_lengths_y, offset_dist_y)
x2_quake = shear_report.carry_list(quake_loads, floor_lengths_x, overhang_dist_x, distances_x, shear_layer_lengths_x, offset_dist_x)
y2_quake = shear_report.carry_list(quake_loads, floor_lengths_y, overhang_dist_y, distances_y, shear_layer_lengths_y, offset_dist_y)

# Here we create lists with the location-based loads and text showing the work for each force and direction
[windx, windx_text] = shear_report.load_add(x2_wind, x1_wind)
[windy, windy_text] = shear_report.load_add(y2_wind, y1_wind)
[quakex, quakex_text] = shear_report.load_add(x2_quake, x1_quake)
[quakey, quakey_text] = shear_report.load_add(y2_quake, y1_quake)

# Here we create the strings showing the full calculations for each force and direction
wind_x = shear_report.shear_text_direction(windx, windx_text, shear_layers_x, x_wind, floor_lengths_x, distances_x, 'x', overhang_dist_x, shear_layer_lengths_x)
quake_x = shear_report.shear_text_direction(quakex, quakex_text, shear_layers_x, quake_loads, floor_lengths_x, distances_x, 'x', overhang_dist_x, shear_layer_lengths_x)
wind_y = shear_report.shear_text_direction(windy, windy_text, shear_layers_y, y_wind, floor_lengths_y, distances_y, 'y', overhang_dist_y, shear_layer_lengths_y)
quake_y = shear_report.shear_text_direction(quakey, quakey_text, shear_layers_y, quake_loads, floor_lengths_y, distances_y, 'y', overhang_dist_y, shear_layer_lengths_y)
shear_complete = shear_report.printer(wind_x, wind_y, quake_x, quake_y)

# SHEAR WALL CALCULATIONS END, WORD DOCUMENT CREATING BEGINS

from docx import Document

project_name = project_name + '.docx'

document = Document()

document.add_paragraph(wind_complete)

document.add_page_break()

document.add_paragraph(quake_complete)

document.add_page_break()

document.add_paragraph(shear_complete)

document.save(save_location + project_name)